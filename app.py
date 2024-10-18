import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt
from groq import Groq
import re

# Load the API key from Streamlit secrets
api_key = st.secrets["GROQ_API_KEY"]

client = Groq(api_key=api_key)
my_llm = "llama-3.2-90b-text-preview"

# Function to generate the lesson plan using AI
def generate_lesson_plan(competency, subject, grade_level, selected_strategies, content, past_lesson, part_a, part_b, part_c, part_d, part_e, part_f, part_g, part_h, part_i, language):
    prompt = f"""
    Generate a lesson plan based on the following parameters:
    Competency: {competency}
    Subject: {subject}
    Grade Level: {grade_level}
    Strategies: {selected_strategies}
    Content: {content}
    past_lesson: {past_lesson}

    Apply the following structure and the {language} language to generate the lesson plan using plain and simple words and in human-like tone:
    Be sure to use level 2 heading and bulleted list always.

    Integration: Enumerate English, Math and Science integration to the subject if any.

    Use the dollowing format:

    Header 3: A. Reviewing or Presenting the New Lesson. Time Limit: {part_a} minutes.
    1. what will the teacher do here? 
    2. HOTS questions to connect past lessons with the new topic, encouraging students to recall relevant information.
        
    Header 3: . Establishing a Purpose for the Lesson. Time Limit: {part_b} minutes.
    1. what will the teacher do here?  
    2. 2 HOTS questions or relevant real-world scenarios that relates to the competency being addressed.
        
    Header 3: C. Presenting Examples/Instances of the New Lesson. Time Limit: {part_c} minutes.
    1. what will the teacher do here? 
    2. List procedure of doing the task in brief anc concise manner. Utilize multimedia resources or real-life demonstrations that incorporate 21st-century skills, such as critical thinking and collaboration.
    The teacher will use appropriate strategy: {selected_strategies}.

    Header 3: D. Discussing New Concepts and Practicing New Skills #1. Time Limit: {part_d} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner considering {selected_strategies} to facilitate understanding.
    
    Header 3: E. Discussing New Concepts and Practicing New Skills #2. Time Limit: {part_e} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner using the appropriate strategy: {selected_strategies}.
      
    
    Header 3: F. Developing Mastery. Time Limit: {part_f} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner using {selected_strategies}.
        
 
    Header 3: G. Finding Practical Applications of Concepts. Time Limit: {part_g} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner considering concepts learned that can be applied in their lives or communities.
  
   Header 3: H. Making Generalizations and Abstractions about the Lesson. Time Limit: {part_h} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner.
      

   Header 3: I. Evaluating Learning. Time Limit: {part_i} minutes.
    1. what will the teacher do here? 
    2. List the procedure of doing the task in brief anc concise manner.The teacher will assign a quiz or project that assesses understanding and application of the lesson content.
    The teacher will use formative assessments like exit tickets where students reflect on what they learned. """
    

    
    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model=my_llm
    )
    return chat_completion.choices[0].message.content

# Function to format the generated lesson plan
def format_lesson_plan(lesson_plan_data):
    formatted_plan = "Daily Lesson Log\n\n"
    sections = lesson_plan_data.split('\n\n')

    for section in sections:
        lines = section.split('\n')
        if lines[0].strip().startswith(('Integration', 'A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.', 'H.', 'I.')):
            formatted_plan += f"**{lines[0].strip()}**\n\n"
            formatted_plan += '\n'.join(lines[1:]).strip() + '\n\n'
        else:
            formatted_plan += section.strip() + '\n\n'

    # Remove extra asterisks from key-value pairs
    formatted_plan = '\n'.join([line.replace('**: **', ': ') if ': ' in line else line for line in formatted_plan.split('\n')])

    return formatted_plan

# Function to export the lesson plan to DOCX
def export_to_docx(lesson_plan, raw_lesson_plan):
    doc = Document()
    title = lesson_plan.split('\n')[0].strip()
    doc.add_heading(title, level=1)
    
    lines = lesson_plan.split('\n')[1:]
    in_list = False
    for line in lines:
        line = line.strip()
        if line.startswith('**') and line.endswith('**'):
            current_section = line.strip('**')
            p = doc.add_paragraph()
            run = p.add_run(current_section)
            run.bold = True
        elif line.startswith('-**') or line.startswith('- **'):
            p = doc.add_paragraph(line.lstrip('-**').strip(), style='List Bullet')
            p.runs[0].bold = True
        elif line.startswith('-'):
            # Strip leading dashes and whitespace
            bullet_text = line.lstrip('-').strip()
            if bullet_text:  # Only add if there's actual text
               doc.add_paragraph(bullet_text, style='List Bullet')
        elif ':' in line:
            key, value = line.split(':', 1)
            p = doc.add_paragraph()
            p.add_run(key.strip('** ')).bold = True
            p.add_run(f": {value.strip('** ')}")
        elif line:
            p = doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Raw AI-Generated Version", level=1)
    doc.add_paragraph(raw_lesson_plan)
    
    doc_file = BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)
    return doc_file

#streamlit styling
st.markdown("""
<style>
.stTextInput > label {
    color: #e516a4d6; /* Change label color */
}
.stMultiSelect > label {
    color: #e516a4d6; /* Change label color */
}    
</style>
""", unsafe_allow_html=True)

   
# Streamlit app layout
st.title("📚 DLL Generator with AI")
st.caption(f"Generated using {my_llm}. Developed by ebb with AI assistance.")

# Organizing input sections into two columns
col1, col2 = st.columns(2)

with col1:
    # we use markdown (**)here to make the label bold
    language = st.text_input("**Language**:","english")
    competency = st.text_input("**Competency:**", "required")
    subject = st.text_input("**Subject:**", "CSS")
    grade_level = st.text_input("**Grade Level:**", "11")

with col2:
    strategies = ["Project-Based Learning", "Collaborative Learning", "Real-World Applications", "Technology Integration", "Differentiated Instruction"]
    selected_strategies = st.multiselect("**Teaching Strategies:**", strategies)
    content = st.text_area("**Content:**", "required")
    past_lesson = st.text_input("**Past Lesson:**", "required")

st.markdown("---")  # Horizontal divider for clarity

# Two-column layout for time limits
col1, col2 = st.columns(2)
with col1:
    part_a = st.text_input("**A. Reviewing or Presenting the New Lesson (minutes):**", "10")
    part_c = st.text_input("**C. Presenting Examples of the New Lesson (minutes):**", "10")
    part_e = st.text_input("**E. Discussing New Skills #2 (minutes):**", "10")
    part_g = st.text_input("**G. Finding Practical Applications (minutes):**", "20")
    part_i = st.text_input("**I. Evaluating Learning (minutes):**", "20")

with col2:
    part_b = st.text_input("**B. Establishing a Purpose (minutes):**", "10")
    part_d = st.text_input("**D. Discussing New Skills #1 (minutes):**", "10")
    part_f = st.text_input("**F. Developing Mastery (minutes):**", "20")
    part_h = st.text_input("**H. Making Generalizations (minutes)**:", "10")

if st.button("Generate Lesson Plan"):
    if language and competency and subject and grade_level and selected_strategies and content and past_lesson:
        raw_lesson_plan = generate_lesson_plan(
            competency, subject, grade_level, selected_strategies, content, past_lesson, 
            part_a, part_b, part_c, part_d, part_e, part_f, part_g, part_h, part_i, language)
        
        formatted_lesson_plan = format_lesson_plan(raw_lesson_plan)

        st.markdown(formatted_lesson_plan)

        docx_file = export_to_docx(formatted_lesson_plan, raw_lesson_plan)

        st.download_button(
                label="Download Lesson Plan (DOCX with Raw AI Output)",
                data=docx_file,
                file_name="lesson_plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    else:
        st.error("Please fill in all required fields!")
